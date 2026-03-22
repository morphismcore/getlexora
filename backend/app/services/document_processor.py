"""
Belge isleme servisi.
PDF, DOCX dosyalarindan metin cikarir, yapilandirir.
AI kullanmaz -- pure extraction.
"""

import io
import re

import fitz  # PyMuPDF
from docx import Document

from app.services.citation_verifier import PATTERNS


class DocumentProcessor:
    """PDF ve DOCX belgelerinden metin cikarma ve analiz."""

    # ── Belge turu tespiti icin anahtar kelimeler ─────────────────────

    _TYPE_KEYWORDS: dict[str, list[str]] = {
        "dilekce": [
            "SAYIN MAHKEME",
            "ARZ VE TALEP EDERİZ",
            "ARZ EDERIM",
            "SAYIN HAKİMLİĞİ",
            "SAYIN BAŞKANLIĞI",
            "DAVACI",
            "DAVALI",
            "DAVA DİLEKÇESİ",
            "CEVAP DİLEKÇESİ",
        ],
        "karar": [
            "KARAR",
            "HÜKÜM",
            "T.C.",
            "DAİRE",
            "DAVA DOSYASININ",
            "GEREĞİ DÜŞÜNÜLDÜ",
            "SONUÇ",
            "YARGITAY",
            "DANIŞTAY",
        ],
        "bilirkisi_raporu": [
            "BİLİRKİŞİ RAPORU",
            "BİLİRKİŞİ",
            "RAPOR",
            "UZMAN GÖRÜŞÜ",
            "İNCELEME RAPORU",
        ],
        "sozlesme": [
            "SÖZLEŞME",
            "TARAFLAR",
            "MADDE 1",
            "SÖZLEŞMENİN KONUSU",
            "YÜKÜMLÜLÜKLER",
            "İMZA",
        ],
        "ihtarname": [
            "İHTARNAME",
            "İHTAR EDEN",
            "MUHATAP",
            "İHTAR OLUNAN",
            "NOTER",
        ],
    }

    # ── Regex kaliplari ───────────────────────────────────────────────

    _RE_DAVACI = re.compile(
        r"DAVACI(?:LAR)?\s*[:\-]?\s*(.+?)(?:\n|DAVALI|VEKİL)",
        re.IGNORECASE | re.DOTALL,
    )
    _RE_DAVALI = re.compile(
        r"DAVALI(?:LAR)?\s*[:\-]?\s*(.+?)(?:\n|VEKİL|DAVA|KONU)",
        re.IGNORECASE | re.DOTALL,
    )
    _RE_VEKIL = re.compile(
        r"VEKİL(?:İ|LERİ)?\s*[:\-]?\s*(.+?)(?:\n|DAVA|KONU|DAVACI|DAVALI)",
        re.IGNORECASE | re.DOTALL,
    )

    _RE_MAHKEME = re.compile(
        r"((?:T\.?C\.?\s+)?(?:\w+\s+){0,5}(?:MAHKEMESİ|MAHKEMESI|Mahkemesi|HAKİMLİĞİ|Hakimliği|BAŞKANLIĞI))",
        re.IGNORECASE,
    )
    _RE_ESAS = re.compile(
        r"(?:ESAS\s*(?:NO|NUMARASI)?\s*[:\-]?\s*|E\.\s*)(\d{4}\s*/\s*\d+)",
        re.IGNORECASE,
    )
    _RE_KARAR = re.compile(
        r"(?:KARAR\s*(?:NO|NUMARASI)?\s*[:\-]?\s*|K\.\s*)(\d{4}\s*/\s*\d+)",
        re.IGNORECASE,
    )
    _RE_TARIH = re.compile(
        r"(\d{1,2}[./]\d{1,2}[./]\d{4})",
    )

    # ── PDF ───────────────────────────────────────────────────────────

    def extract_text_from_pdf(self, file_bytes: bytes) -> dict:
        """
        PDF'den metin cikar.
        Returns: {
            "text": "tam metin",
            "pages": sayfa_sayisi,
            "page_texts": ["sayfa 1 metni", ...],
            "metadata": {"title": "...", "author": "..."}
        }
        """
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_texts: list[str] = []

        for page in doc:
            page_texts.append(page.get_text())

        metadata_raw = doc.metadata or {}
        metadata = {
            "title": metadata_raw.get("title", "") or "",
            "author": metadata_raw.get("author", "") or "",
            "subject": metadata_raw.get("subject", "") or "",
            "creator": metadata_raw.get("creator", "") or "",
            "producer": metadata_raw.get("producer", "") or "",
        }

        full_text = "\n\n".join(page_texts)
        page_count = len(page_texts)
        doc.close()

        return {
            "text": full_text,
            "pages": page_count,
            "page_texts": page_texts,
            "metadata": metadata,
        }

    # ── DOCX ──────────────────────────────────────────────────────────

    def extract_text_from_docx(self, file_bytes: bytes) -> dict:
        """
        DOCX'ten metin cikar.
        Returns: {
            "text": "tam metin",
            "paragraphs": sayi,
            "metadata": {"title": "...", "author": "..."}
        }
        """
        doc = Document(io.BytesIO(file_bytes))

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Tablolardan metin cikar
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n--- Tablolar ---\n" + "\n".join(table_texts)

        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
        }

        return {
            "text": full_text,
            "paragraphs": len(paragraphs),
            "metadata": metadata,
        }

    # ── Belge turu tespiti ────────────────────────────────────────────

    def detect_document_type(self, text: str) -> str:
        """
        Belge turunu tespit et (regex/keyword bazli):
        - dilekce, karar, bilirkisi_raporu, sozlesme, ihtarname, diger
        """
        upper_text = text.upper()
        scores: dict[str, int] = {}

        for doc_type, keywords in self._TYPE_KEYWORDS.items():
            score = 0
            for kw in keywords:
                if kw.upper() in upper_text:
                    score += 1
            scores[doc_type] = score

        if not scores:
            return "diger"

        best_type = max(scores, key=lambda k: scores[k])
        if scores[best_type] >= 2:
            return best_type

        return "diger"

    # ── Taraf cikarma ─────────────────────────────────────────────────

    def extract_parties(self, text: str) -> dict:
        """
        Taraflari cikar (regex bazli):
        - Davaci/Davacilar
        - Davali/Davalilar
        - Vekil/Vekilleri
        """
        result: dict[str, str | None] = {
            "davaci": None,
            "davali": None,
            "davaci_vekili": None,
            "davali_vekili": None,
        }

        m = self._RE_DAVACI.search(text)
        if m:
            result["davaci"] = m.group(1).strip()[:200]

        m = self._RE_DAVALI.search(text)
        if m:
            result["davali"] = m.group(1).strip()[:200]

        # Vekil - birden fazla olabilir
        vekiller = self._RE_VEKIL.findall(text)
        if vekiller:
            result["davaci_vekili"] = vekiller[0].strip()[:200]
            if len(vekiller) > 1:
                result["davali_vekili"] = vekiller[1].strip()[:200]

        return result

    # ── Dava bilgileri cikarma ────────────────────────────────────────

    def extract_case_info(self, text: str) -> dict:
        """
        Dava bilgilerini cikar:
        - Mahkeme adi
        - Esas numarasi
        - Karar numarasi
        - Tarih
        """
        result: dict[str, str | None] = {
            "mahkeme": None,
            "esas_no": None,
            "karar_no": None,
            "tarih": None,
        }

        m = self._RE_MAHKEME.search(text)
        if m:
            result["mahkeme"] = m.group(1).strip()

        m = self._RE_ESAS.search(text)
        if m:
            result["esas_no"] = m.group(1).strip().replace(" ", "")

        m = self._RE_KARAR.search(text)
        if m:
            result["karar_no"] = m.group(1).strip().replace(" ", "")

        m = self._RE_TARIH.search(text)
        if m:
            result["tarih"] = m.group(1)

        return result

    # ── Referans cikarma ──────────────────────────────────────────────

    def extract_citations_from_document(self, text: str) -> list[dict]:
        """Belgeden tum hukuki referanslari cikar (citation_verifier regex'lerini kullan)."""
        citations: list[dict] = []
        seen: set[str] = set()

        for pattern_name, pattern in PATTERNS.items():
            for m in pattern.finditer(text):
                raw = m.group(0)
                if raw in seen:
                    continue
                seen.add(raw)
                citations.append({
                    "raw_text": raw,
                    "pattern_type": pattern_name,
                })

        return citations
