"""
Legal document export — DOCX and PDF generation.
Takes structured document data and produces properly formatted files.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import io
import string


class DocumentExportService:
    """Generates DOCX and PDF files from structured legal document data."""

    # ------------------------------------------------------------------ #
    #  DOCX generation
    # ------------------------------------------------------------------ #

    def generate_docx(self, doc_data: dict) -> bytes:
        """
        Generate a properly formatted DOCX from structured document data.

        doc_data format:
        {
            "docType": "dava_dilekce",
            "header": {
                "mahkeme": "Istanbul ( ). Is Mahkemesi'ne",
                "davaci": "Ahmet Yilmaz",
                "davaci_tc": "12345678901",
                "davaci_adres": "...",
                "davaci_vekili": "Av. Mehmet ...",
                "davali": "XYZ Ltd. Sti.",
                "davali_adres": "...",
                "konu": "Feshin gecersizligi..."
            },
            "blocks": [ ... ]
        }
        """
        doc = Document()

        # -- Page setup: A4, 2.5 cm margins --
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # -- Default style tweaks --
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.15

        header = doc_data.get("header", {})
        blocks = doc_data.get("blocks", [])

        # -- Mahkeme header --
        if header.get("mahkeme"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header["mahkeme"])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

        doc.add_paragraph()  # blank line

        # -- Party information --
        party_fields = [
            ("DAVACI", "davaci"),
            ("T.C. KİMLİK NO", "davaci_tc"),
            ("ADRES", "davaci_adres"),
            ("VEKİLİ", "davaci_vekili"),
            ("DAVALI", "davali"),
            ("ADRES", "davali_adres"),
            ("KONU", "konu"),
        ]
        for label, key in party_fields:
            value = header.get(key)
            if value:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_label = p.add_run(f"{label}\t: ")
                run_label.bold = True
                run_label.font.size = Pt(11)
                run_label.font.name = "Times New Roman"
                run_value = p.add_run(value)
                run_value.font.size = Pt(11)
                run_value.font.name = "Times New Roman"

        doc.add_paragraph()  # blank line

        # -- Blocks --
        numbered_counter = 0
        for block in blocks:
            btype = block.get("type", "free_text")
            content = block.get("content", "")

            if btype == "section_header":
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(content.upper())
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

            elif btype == "numbered_paragraph":
                numbered_counter += 1
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(f"{numbered_counter}. {content}")
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

                # Sub-paragraphs (a, b, c, ...)
                children = block.get("children", [])
                for idx, child in enumerate(children):
                    letter = string.ascii_lowercase[idx] if idx < 26 else str(idx + 1)
                    sp = doc.add_paragraph()
                    sp.paragraph_format.left_indent = Cm(1.25)
                    sp.paragraph_format.line_spacing = 1.15
                    run_sub = sp.add_run(f"{letter}) {child.get('content', '')}")
                    run_sub.font.size = Pt(11)
                    run_sub.font.name = "Times New Roman"

            elif btype == "evidence_item":
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(content)
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

            else:
                # free_text or any unknown type
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(content)
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

        # -- Footer: date + signature --
        doc.add_paragraph()
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_date = p_date.add_run("Tarih: …/…/……")
        run_date.font.size = Pt(11)
        run_date.font.name = "Times New Roman"

        p_sign = doc.add_paragraph()
        p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if header.get("davaci_vekili"):
            sign_text = header["davaci_vekili"]
        elif header.get("davaci"):
            sign_text = header["davaci"]
        else:
            sign_text = "İmza"
        run_sign = p_sign.add_run(sign_text)
        run_sign.font.size = Pt(11)
        run_sign.font.name = "Times New Roman"

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------ #
    #  PDF generation
    # ------------------------------------------------------------------ #

    def generate_pdf(self, doc_data: dict) -> bytes:
        """
        Generate PDF from structured document data.
        Uses fpdf2 with built-in unicode support for Turkish characters.
        """
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=25)
        pdf.add_page()
        pdf.set_margins(25, 25, 25)

        # Try to use DejaVu if available, fall back to Helvetica
        _font_family = self._setup_pdf_font(pdf)
        st = lambda text: self._safe_text(text, _font_family)

        header = doc_data.get("header", {})
        blocks = doc_data.get("blocks", [])

        # -- Mahkeme header --
        if header.get("mahkeme"):
            pdf.set_font(_font_family, "B", 12)
            pdf.cell(0, 8, st(header["mahkeme"]), new_x="LMARGIN", new_y="NEXT", align="C")

        pdf.ln(6)

        # -- Party information --
        party_fields = [
            ("DAVACI", "davaci"),
            ("T.C. KIMLIK NO", "davaci_tc"),
            ("ADRES", "davaci_adres"),
            ("VEKILI", "davaci_vekili"),
            ("DAVALI", "davali"),
            ("ADRES", "davali_adres"),
            ("KONU", "konu"),
        ]
        pdf.set_font(_font_family, "", 11)
        for label, key in party_fields:
            value = header.get(key)
            if value:
                pdf.set_font(_font_family, "B", 11)
                pdf.cell(40, 6, st(label), new_x="RIGHT", new_y="TOP")
                pdf.set_font(_font_family, "", 11)
                pdf.cell(5, 6, ": ", new_x="RIGHT", new_y="TOP")
                pdf.multi_cell(0, 6, st(value), new_x="LMARGIN", new_y="NEXT")

        pdf.ln(6)

        # -- Blocks --
        numbered_counter = 0
        for block in blocks:
            btype = block.get("type", "free_text")
            content = block.get("content", "")

            if btype == "section_header":
                pdf.ln(4)
                pdf.set_font(_font_family, "B", 11)
                pdf.cell(0, 7, st(content.upper()), new_x="LMARGIN", new_y="NEXT")

            elif btype == "numbered_paragraph":
                numbered_counter += 1
                pdf.set_font(_font_family, "", 11)
                text = f"{numbered_counter}. {content}"
                pdf.multi_cell(0, 6, st(f"    {text}"), new_x="LMARGIN", new_y="NEXT")

                children = block.get("children", [])
                for idx, child in enumerate(children):
                    letter = string.ascii_lowercase[idx] if idx < 26 else str(idx + 1)
                    pdf.set_font(_font_family, "", 11)
                    child_text = f"        {letter}) {child.get('content', '')}"
                    pdf.multi_cell(0, 6, st(child_text), new_x="LMARGIN", new_y="NEXT")

            elif btype == "evidence_item":
                pdf.set_font(_font_family, "", 11)
                pdf.cell(8, 6, "-", new_x="RIGHT", new_y="TOP")
                pdf.multi_cell(0, 6, st(content), new_x="LMARGIN", new_y="NEXT")

            else:
                pdf.set_font(_font_family, "", 11)
                pdf.multi_cell(0, 6, st(content), new_x="LMARGIN", new_y="NEXT")

        # -- Footer: date + signature --
        pdf.ln(10)
        pdf.set_font(_font_family, "", 11)
        pdf.cell(0, 6, "Tarih: .../.../..........", new_x="LMARGIN", new_y="NEXT", align="R")

        if header.get("davaci_vekili"):
            sign_text = header["davaci_vekili"]
        elif header.get("davaci"):
            sign_text = header["davaci"]
        else:
            sign_text = "Imza"
        pdf.cell(0, 6, st(sign_text), new_x="LMARGIN", new_y="NEXT", align="R")

        return pdf.output()

    @staticmethod
    def _safe_text(text: str, font_family: str) -> str:
        """Make text safe for non-Unicode fonts by replacing Turkish chars."""
        if font_family != "Helvetica":
            return text
        # Helvetica can't handle Turkish special chars
        replacements = {
            "İ": "I", "ı": "i", "Ğ": "G", "ğ": "g",
            "Ü": "U", "ü": "u", "Ş": "S", "ş": "s",
            "Ö": "O", "ö": "o", "Ç": "C", "ç": "c",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    @staticmethod
    def _setup_pdf_font(pdf: FPDF) -> str:
        """Set up a font with Turkish character support. Returns the font family name."""
        import os
        import glob

        # Search common font locations
        search_patterns = [
            "/usr/share/fonts/**/DejaVuSans.ttf",
            "/usr/share/fonts-dejavu-core/DejaVuSans.ttf",
            "/usr/local/share/fonts/**/DejaVuSans.ttf",
        ]
        bold_patterns = [
            "/usr/share/fonts/**/DejaVuSans-Bold.ttf",
            "/usr/share/fonts-dejavu-core/DejaVuSans-Bold.ttf",
            "/usr/local/share/fonts/**/DejaVuSans-Bold.ttf",
        ]

        font_path = None
        bold_path = None

        for pattern in search_patterns:
            matches = glob.glob(pattern, recursive=True)
            if matches:
                font_path = matches[0]
                break

        for pattern in bold_patterns:
            matches = glob.glob(pattern, recursive=True)
            if matches:
                bold_path = matches[0]
                break

        if font_path:
            pdf.add_font("DejaVu", "", font_path)
            if bold_path:
                pdf.add_font("DejaVu", "B", bold_path)
            else:
                # Use regular as bold fallback
                pdf.add_font("DejaVu", "B", font_path)
            return "DejaVu"

        # Fallback: use fpdf2's built-in Unicode support
        # This works with fpdf2 >= 2.7 which bundles a Unicode font
        try:
            pdf.add_font("NotoSans", "", str(os.path.join(os.path.dirname(__file__), "fonts", "NotoSans-Regular.ttf")))
        except Exception:
            pass

        # Last resort: Helvetica with ASCII-safe text
        return "Helvetica"
